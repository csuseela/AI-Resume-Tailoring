from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


class DocxWriterService:
    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def write(self, company: str, role: str, markdown: str) -> Path:
        doc = Document()
        self._set_margins(doc)
        self._set_default_font(doc)

        lines = markdown.strip().splitlines()
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("# "):
                self._add_name(doc, stripped[2:].strip())
            elif stripped.startswith("**") and stripped.endswith("**"):
                self._add_subtitle(doc, stripped.strip("*").strip())
            elif stripped.startswith("## "):
                self._add_section_heading(doc, stripped[3:].strip())
            elif stripped.startswith("### "):
                self._add_role_heading(doc, stripped[4:].strip())
            elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
                self._add_date_line(doc, stripped.strip("*").strip())
            elif stripped.startswith("- "):
                self._add_bullet(doc, stripped[2:].strip())
            elif stripped == "---":
                self._add_divider(doc)
            elif stripped:
                if "linkedin" in stripped.lower() or "github" in stripped.lower():
                    self._add_contact_line(doc, stripped)
                else:
                    self._add_paragraph(doc, stripped)

        slug_company = re.sub(r"[^a-zA-Z0-9]+", "_", company.strip().lower()).strip("_")
        slug_role = re.sub(r"[^a-zA-Z0-9]+", "_", role.strip().lower()).strip("_")
        date_str = datetime.now().strftime("%Y-%m-%d")
        filename = f"{date_str}_{slug_company}_{slug_role}.docx"
        path = self.output_dir / filename
        doc.save(str(path))
        logger.info("Word document saved: %s", path)
        return path

    @staticmethod
    def _set_margins(doc: Document) -> None:
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    @staticmethod
    def _set_default_font(doc: Document) -> None:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.space_before = Pt(0)

    @staticmethod
    def _add_name(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _add_subtitle(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        p.paragraph_format.space_after = Pt(4)

    @staticmethod
    def _add_contact_line(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        p.paragraph_format.space_after = Pt(6)

    @staticmethod
    def _add_divider(doc: Document) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\u2500" * 80)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _add_section_heading(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        bottom = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "1F4E79")
        bottom.append(b)
        pPr.append(bottom)

    @staticmethod
    def _add_role_heading(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

    @staticmethod
    def _add_date_line(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        p.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _add_bullet(doc: Document, text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.3)

    @staticmethod
    def _add_paragraph(doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
